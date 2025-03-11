import docx
import re
import difflib
import streamlit as st
from fuzzywuzzy import fuzz

def preprocess_question(question):
    match = re.search(r'^(.*?)\?', question)
    if match:
        return match.group(1) + "?"
    return question

def extract_and_match(docx_file, prompt_file, similarity_threshold=90):
    matched_lines = []
    try:
        doc = docx.Document(docx_file)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        lines = full_text.split('\n')

        with open(prompt_file, "r", encoding="utf-8") as file:
            questions = [line.strip() for line in file]

        for question in questions:
            best_match = None
            best_similarity = 0
            processed_question = preprocess_question(question)

            for line in lines:
                similarity = fuzz.ratio(processed_question, line)
                if similarity > best_similarity:
                    best_similarity = similarity
                    best_match = line

            if best_similarity >= similarity_threshold:
                matched_lines.append(best_match)
        
        return matched_lines
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return []

def clean_text(text):
    text = re.sub(r'####\s*\**|\**', '', text)
    text = re.sub(r'[-]{3,}', '', text)
    text = text.replace("\n", " ").strip()
    text = " ".join(text.split())
    return text.lower()

def is_heading(text, matched_headings):
    text_cleaned = clean_text(text)
    for heading in matched_headings:
        heading_cleaned = clean_text(heading)
        similarity = difflib.SequenceMatcher(None, text_cleaned, heading_cleaned).ratio()
        if text_cleaned == heading_cleaned or similarity >= 0.95:
            return True
    return False

def extract_sections_from_docx(file_path, matched_headings):
    try:
        doc = docx.Document(file_path)
        sections = {}
        current_section = None
        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else "Unknown"
            potential_headings = re.split(r'(#### .*?)\n', text)
            for segment in potential_headings:
                segment = segment.strip()
                if not segment:
                    continue
                if is_heading(segment, matched_headings) or "Heading" in style:
                    current_section = segment
                    sections[current_section] = ""
                elif current_section:
                    sections[current_section] += segment + "\n"
        return sections
    except Exception as e:
        st.error(f"Error extracting sections from {file_path}: {e}")
        return {}

def merge_documents(original_file, updated_file, matched_headings):
    original_sections = extract_sections_from_docx(original_file, matched_headings)
    updated_sections = extract_sections_from_docx(updated_file, matched_headings)
    if not original_sections or not updated_sections:
        st.error("Unable to process documents! Check the format and try again.")
        return None
    merged_sections = original_sections.copy()
    for original_heading in original_sections.keys():
        best_match = find_best_match(original_heading, updated_sections.keys())
        if best_match:
            merged_sections[original_heading] = updated_sections[best_match]
    merged_doc = docx.Document()
    for heading, content in merged_sections.items():
        merged_doc.add_paragraph(heading, style='Heading1')
        merged_doc.add_paragraph(content)
    return merged_doc

def find_best_match(original_heading, updated_headings):
    best_match = None
    highest_similarity = 0.0
    for updated_heading in updated_headings:
        similarity = difflib.SequenceMatcher(None, original_heading, updated_heading).ratio()
        if similarity > highest_similarity and similarity >= 0.95:
            highest_similarity = similarity
            best_match = updated_heading
    return best_match

def main():
    st.title("Docx Matcher & Merger")
    st.write("Upload your documents to extract and merge sections based on matched questions.")
    
    uploaded_docx = st.file_uploader("Upload Original DOCX", type=["docx"])
    uploaded_updated_docx = st.file_uploader("Upload Updated DOCX", type=["docx"])
    uploaded_prompt = st.file_uploader("Upload Prompt TXT", type=["txt"])
    
    if st.button("Extract and Merge"):
        if uploaded_docx and uploaded_updated_docx and uploaded_prompt:
            with open("temp_prompt.txt", "wb") as f:
                f.write(uploaded_prompt.getvalue())
            matched_headings = extract_and_match(uploaded_docx, "temp_prompt.txt")
            if matched_headings:
                merged_doc = merge_documents(uploaded_docx, uploaded_updated_docx, matched_headings)
                if merged_doc:
                    merged_doc.save("merged_report.docx")
                    st.success("Merged document created successfully!")
                    with open("merged_report.docx", "rb") as f:
                        st.download_button("Download Merged Report", f, file_name="merged_report.docx")
                else:
                    st.error("Failed to merge documents.")
        else:
            st.error("Please upload all required files.")

if __name__ == "__main__":
    main()
