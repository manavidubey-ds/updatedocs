import re
import difflib
import docx
import tkinter as tk
from tkinter import filedialog, messagebox

# -------------------- LOAD MATCHED HEADINGS --------------------

def load_matched_headings(file_path):
    """Reads matched_questions.txt and returns a set of cleaned headings."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            matched_headings = set(line.strip() for line in f.readlines() if line.strip())
            print("\n📂 Matched Questions from matched_questions.txt:")
            for heading in matched_headings:
                print(f"🔹 '{heading}' (Cleaned: '{clean_text(heading)}')")
            return matched_headings
    except Exception as e:
        print(f"❌ Error reading matched_questions.txt: {e}")
        return set()

# -------------------- CLEAN & MATCH HEADINGS --------------------

def clean_text(text):
    """Removes Markdown-style formatting like ####, **, --- and extra spaces."""
    text = re.sub(r'####\s*\**|\**', '', text)  # Remove '####' and '**'
    text = re.sub(r'[-]{3,}', '', text)  # Remove '---' dividers
    text = text.replace("\n", " ").strip()  # Remove newlines, keep space
    text = " ".join(text.split())  # Normalize spaces
    return text.lower()  # Convert to lowercase for uniform comparison

def is_heading(text, matched_headings):
    """Checks if text matches a heading in matched_questions.txt (exact or 95% similar)."""
    text_cleaned = clean_text(text)

    for heading in matched_headings:
        heading_cleaned = clean_text(heading)
        similarity = difflib.SequenceMatcher(None, text_cleaned, heading_cleaned).ratio()

        if similarity >= 0.90:
            print(f"🔍 Comparing: '{text_cleaned}' ↔ '{heading_cleaned}' (Similarity: {similarity:.2f})")

        if text_cleaned == heading_cleaned or similarity >= 0.95:
            print(f"✅ MATCH FOUND: '{text}' (Similarity: {similarity:.2f})")
            return True

    return False

# -------------------- EXTRACT SECTIONS FROM DOCX --------------------

def extract_sections_from_docx(file_path, matched_headings):
    """
    Extracts sections using matched_questions.txt for heading identification.
    Handles both:
      - Proper Word headings (Heading 1, Heading 2)
      - Markdown-style headings inside paragraphs (#### 1. ...)
    """
    try:
        doc = docx.Document(file_path)
        sections = {}
        current_section = None

        print(f"\n📖 Processing {file_path}...")

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            style = para.style.name if para.style else "Unknown"

            # 🛠️ **Fix: Manually split paragraphs on '####' if present**
            potential_headings = re.split(r'(#### .*?)\n', text)

            for segment in potential_headings:
                segment = segment.strip()
                if not segment:
                    continue  # Skip empty segments

                # 🔍 **Detect true headings (styled or matched from txt)**
                if is_heading(segment, matched_headings) or "Heading" in style:
                    print(f"✅ Found Heading: '{segment}' (Style: '{style}')")
                    current_section = segment
                    sections[current_section] = ""
                elif current_section:
                    sections[current_section] += segment + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if is_heading(cell_text, matched_headings):
                        print(f"✅ Found Heading in Table: '{cell_text}'")
                        current_section = cell_text
                        sections[current_section] = ""

        print(f"🔍 Extracted {len(sections)} headings from {file_path}")
        return sections

    except Exception as e:
        print(f"❌ Error extracting sections from {file_path}: {e}")
        return {}

# -------------------- MERGE DOCUMENTS --------------------

def merge_documents(original_file, updated_file, matched_headings, output_file="merged_report.docx"):
    """
    Merges the original and updated documents by replacing sections that match `matched_questions.txt` headings.
    """
    print("\n🔄 Extracting headings from original document...")
    original_sections = extract_sections_from_docx(original_file, matched_headings)

    print("🔄 Extracting headings from updated document...")
    updated_sections = extract_sections_from_docx(updated_file, matched_headings)

    if not original_sections or not updated_sections:
        messagebox.showerror("Error", "❌ Unable to process documents! Check the format and try again.")
        return

    merged_sections = original_sections.copy()

    for original_heading in original_sections.keys():
        best_match = find_best_match(original_heading, updated_sections.keys())
        if best_match:
            print(f"✅ Replacing section: {original_heading}")
            merged_sections[original_heading] = updated_sections[best_match]  # Replace section

    # Create merged document
    merged_doc = docx.Document()
    for heading, content in merged_sections.items():
        merged_doc.add_paragraph(heading, style='Heading1')
        merged_doc.add_paragraph(content)

    merged_doc.save(output_file)
    print(f"\n🎉 Merging Complete! File saved as: {output_file}")
    messagebox.showinfo("Success", f"Merged document saved as {output_file}")

# -------------------- FIND BEST MATCH FOR HEADINGS --------------------

def find_best_match(original_heading, updated_headings):
    """
    Finds the best matching heading from the updated document with at least 95% similarity.
    """
    best_match = None
    highest_similarity = 0.0

    for updated_heading in updated_headings:
        similarity = difflib.SequenceMatcher(None, original_heading, updated_heading).ratio()
        if similarity > highest_similarity and similarity >= 0.95:  # 95% similarity threshold
            highest_similarity = similarity
            best_match = updated_heading

    return best_match

# -------------------- GUI FOR FILE SELECTION --------------------

def select_file(title):
    """Opens file dialog to select a .docx file and returns the file path."""
    root = tk.Tk()
    root.withdraw()  # Hide main window
    file_path = filedialog.askopenfilename(title=title, filetypes=[("Word Documents", "*.docx")])
    return file_path

# -------------------- MAIN EXECUTION --------------------

if __name__ == "__main__":
    matched_headings = load_matched_headings("matched_questions.txt")

    if not matched_headings:
        print("⚠️ No matched headings found. Check matched_questions.txt.")
    else:
        print("\n📂 Select the ORIGINAL document")
        original_file = select_file("Select the Original Document")

        print("\n📂 Select the UPDATED document")
        updated_file = select_file("Select the Updated Document")

        if original_file and updated_file:
            merge_documents(original_file, updated_file, matched_headings)
        else:
            print("⚠️ Process canceled. Both documents must be selected.")
